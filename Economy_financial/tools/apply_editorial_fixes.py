"""Apply the three editorial corrections from the repository consistency audit.

1. Correct the per-seed order-book event-count range in Section 5.7 to the
   stored range (about 256,000-298,000).
2. Repaint Figure 2's embedded title and footnote so they state that the
   figure covers all four horizons; the plotted data are untouched.
3. Cite Parguel, Benoit-Moreau, and Larceneux (2011) in the Section 1
   greenwashing discussion, resolving the previously uncited reference entry.

No model logic, parameter, result, manifest, or raw output is modified, and
no simulation is executed. The pre-edit manuscript is preserved as a backup.
"""

from __future__ import annotations

import shutil
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont

SOURCE = Path("Francesco_Romano_EU_Environmental_Claims_ABM_Paper_revised.docx")
BACKUP = Path(
    "Francesco_Romano_EU_Environmental_Claims_ABM_Paper_revised"
    ".pre_editorial_fixes_backup.docx")
FONT_PATH = Path("C:/Windows/Fonts/arial.ttf")

FIGURE2_TITLE = "Figure 2. Default-weight winner shares across horizons"
FIGURE2_FOOTNOTE = (
    "Horizons 120, 365, 1,000, and 2,000 days; 200 LHS draws; 3 paired "
    "replications; sampled social discount rate. Simulation-based policy "
    "experiment; not an empirical forecast.")

OLD_RANGE = "about 270,000-300,000 events per seed"
NEW_RANGE = "about 256,000-298,000 events per seed"

PARGUEL_ANCHOR = ("disclose less overall (Lyon and Maxwell, 2011). "
                  "Selective disclosure")
PARGUEL_REPLACEMENT = (
    "disclose less overall (Lyon and Maxwell, 2011). Field evidence on "
    "third-party sustainability ratings similarly indicates that independent "
    "evaluation can deter greenwashing (Parguel, Benoit-Moreau, and "
    "Larceneux, 2011). Selective disclosure")


def find_paragraph(document: Document, contains: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if contains in paragraph.text:
            return paragraph
    raise ValueError(f"paragraph not found: {contains!r}")


def replace_text(paragraph: Paragraph, old: str, new: str) -> None:
    text = paragraph.text
    if old not in text:
        raise ValueError(f"text not found in paragraph: {old!r}")
    updated = text.replace(old, new)
    paragraph.clear()
    paragraph.add_run(updated)


def repaint_figure2(document: Document) -> None:
    for relationship in document.part.rels.values():
        if relationship.reltype != RT.IMAGE:
            continue
        image_part = relationship.target_part
        if str(image_part.partname) != "/word/media/image2.png":
            continue
        image = Image.open(BytesIO(image_part.blob)).convert("RGB")
        draw = ImageDraw.Draw(image)

        # Title band: identical geometry to tools/revise_manuscript.py.
        draw.rectangle((0, 0, image.width, round(image.height * 0.073)),
                       fill="white")
        title_font = ImageFont.truetype(str(FONT_PATH),
                                        round(image.height * 0.037))
        box = draw.textbbox((0, 0), FIGURE2_TITLE, font=title_font)
        x = (image.width - (box[2] - box[0])) / 2
        draw.text((x, round(image.height * 0.024)), FIGURE2_TITLE,
                  fill="black", font=title_font)

        # Footnote band: bottom strip, left-aligned, font auto-fit to width.
        band_top = round(image.height * 0.955)
        draw.rectangle((0, band_top, image.width, image.height), fill="white")
        size = 22
        while size > 12:
            footnote_font = ImageFont.truetype(str(FONT_PATH), size)
            box = draw.textbbox((0, 0), FIGURE2_FOOTNOTE, font=footnote_font)
            if box[2] - box[0] <= image.width - 40:
                break
            size -= 1
        draw.text((20, round(image.height * 0.965)), FIGURE2_FOOTNOTE,
                  fill="black", font=footnote_font)

        output = BytesIO()
        image.save(output, format="PNG")
        image_part._blob = output.getvalue()
        return
    raise ValueError("embedded image /word/media/image2.png not found")


def main() -> None:
    if not BACKUP.exists():
        shutil.copyfile(SOURCE, BACKUP)

    document = Document(str(SOURCE))

    replace_text(find_paragraph(document, OLD_RANGE), OLD_RANGE, NEW_RANGE)
    replace_text(find_paragraph(document, PARGUEL_ANCHOR),
                 PARGUEL_ANCHOR, PARGUEL_REPLACEMENT)
    repaint_figure2(document)

    document.save(str(SOURCE))
    print(f"edited manuscript saved: {SOURCE}")
    print(f"pre-edit backup: {BACKUP}")


if __name__ == "__main__":
    main()
