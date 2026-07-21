"""Strategy-A publication revision: market layer reframed as response, not discipline.

Reads the current revised manuscript, backs it up once, and applies targeted
editorial changes that (i) replace the superseded single-path financial
diagnostic with the completed 30-seed validation campaign, including the
non-reproduced positive-price-impact result, (ii) report the executed
15-replication robustness extension, and (iii) state explicitly that firm
environmental decisions contain no own-share-price feedback, so no policy
conclusion relies on stock-price discipline.  No simulation output is read,
written, or modified.  Every number inserted here is copied from
results/financial_validation/pilot_30seeds/diagnostics/ and
results/replication_robustness/replication_robustness_summary.json.
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


SOURCE = Path("Francesco_Romano_EU_Environmental_Claims_ABM_Paper_revised.docx")
BACKUP = Path(
    "Francesco_Romano_EU_Environmental_Claims_ABM_Paper_revised"
    ".pre_market_reframe_backup.docx")


def find_paragraph(document: Document, startswith: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(startswith):
            return paragraph
    raise ValueError(f"paragraph not found: {startswith!r}")


def replace(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def insert_before(reference: Paragraph, text: str,
                  style: str = "Normal") -> Paragraph:
    element = OxmlElement("w:p")
    reference._p.addprevious(element)
    paragraph = Paragraph(element, reference._parent)
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def insert_table_before(document: Document, reference: Paragraph,
                        rows: list[list[str]]) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        row_properties = table.rows[row_index]._tr.get_or_add_trPr()
        row_properties.append(OxmlElement("w:cantSplit"))
        if row_index == 0:
            row_properties.append(OxmlElement("w:tblHeader"))
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.text = value
            if row_index == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "1F4E78")
                cell._tc.get_or_add_tcPr().append(shading)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    color = OxmlElement("w:color")
                    color.set(qn("w:val"), "FFFFFF")
                    run._r.get_or_add_rPr().append(color)
    reference._p.addprevious(table._tbl)


ABSTRACT = (
    "Environmental-claim governance combines an information problem with an "
    "institutional-capacity problem: regulators observe claims and uncertain "
    "evidence rather than firms' latent environmental states, while preventive "
    "scrutiny may suppress truthful communication. We develop an "
    "information-safe agent-based model linking corporate environmental "
    "claims, consumer demand, investor valuation, workforce responses, a "
    "limit-order-book market, and capacity-constrained supervision. We compare "
    "current EU-style ex-post supervision with two experimental institutions: "
    "a voluntary algorithmic SME pre-screening hub and a certified green-data "
    "connector. A 28-parameter Latin-hypercube design evaluates 200 "
    "configurations at 120, 365, 1,000, and 2,000 days, with three paired "
    "common-random-number replications per configuration; a separate executed "
    "extension re-evaluates 12 representative configurations at 15 "
    "replications. Within the model, the hub reduces exposure-weighted claim "
    "severity in 85.0-90.5% of configurations across horizons, but "
    "systematically increases greenhushing and public and firm compliance "
    "costs. The connector is weak and heterogeneous at 120 days, improves "
    "severity in 60.5% of configurations at 365 days, and becomes robust at "
    "1,000 and 2,000 days (97.0% and 100.0%), while remaining cost-intensive "
    "and increasing evidence-conflict workload. Default rankings reverse "
    "across horizons in 80.5% of matched configurations, showing why policy "
    "comparisons must be conditional on time, evidence quality, capacity, "
    "participation, and normative weights. A 30-seed stylized-facts validation "
    "of the market engine reproduces volatility clustering, positive spread "
    "and depth, trade-sign persistence, and cancellation activity; partially "
    "reproduces fat tails, the volume-volatility relation, and weak linear "
    "return autocorrelation; and does not reproduce positive price impact "
    "under the implemented diagnostic. The market layer is therefore reported "
    "as endogenous market response and financing context: firm greenwashing, "
    "disclosure, and transition-investment rules contain no own-share-price "
    "feedback, and no conclusion relies on stock-price discipline of corporate "
    "behavior. The results are transparent simulation-based policy orderings, "
    "not empirical forecasts or causal estimates.")

MARKET_DIRECTION = (
    "The direction of this financial layer must be stated precisely, because "
    "it bounds what the model can claim. Claims, evidence, and published "
    "supervisory outcomes move investor valuation and hence order flow and "
    "prices, and market prices affect corporate financing proceeds when the "
    "corporate policy sells treasury shares against the order-book midpoint. "
    "A static audit of every identified firm decision method (results/audits/) "
    "finds no other direct or indirect own-share-price input: greenwashing, "
    "disclosure, qualification, and transition-investment rules never read "
    "the firm's own price, order book, or market valuation. The equity market "
    "is therefore an endogenous market-response and financing context, not a "
    "demonstrated price-discipline mechanism, and no result in this paper "
    "relies on stock prices disciplining corporate environmental behavior.")

ROBUSTNESS_DESIGN = (
    "Three paired replications per LHS draw are appropriate for broad global "
    "screening but can under-sample rare escalation, institutional failure, "
    "and tail outcomes. A separate dry-run-first workflow deterministically "
    "selects 12 configurations by maximin coverage of standardized "
    "cross-horizon outcomes and targets 15 paired replications by default "
    "(configurable from 10 to 20). It preserves the authoritative first three "
    "replications and adds only replications 4-15 under the same paired seed "
    "schedule.")

ROBUSTNESS_RESULT = (
    "The extension has been executed. Draws 41, 42, 69, 71, 75, 104, 109, "
    "136, 144, 147, 180, and 187 were extended to 15 paired replications at "
    "each horizon, and all 48 draw-horizon extension files pass validation. "
    "Moving from three to 15 replications narrows the paired 95% interval in "
    "92.3% of the 1,743 draw-metric-policy cells with a finite comparison, "
    "with a median half-width ratio of 0.26. The default-weight winner "
    "computed at three replications is confirmed at 15 replications in 32 of "
    "48 draw-horizon cells (9, 9, 8, and 6 of 12 at 120, 365, 1,000, and "
    "2,000 days), and the complete three-regime ranking is confirmed in 60.4% "
    "of cells. Draw-level winner labels under three replications are "
    "therefore noisy, especially at long horizons; the paper accordingly "
    "emphasizes distributional effects, improvement fractions, and Pareto "
    "membership across the 200 draws rather than draw-level league tables. "
    "The first three replications remain the authoritative campaign, and "
    "extension outputs are stored separately in results/replication_robustness/.")

VALIDATION_INTRO = (
    "The market engine is validated in a dedicated fixed-parameter campaign, "
    "separate from the policy LHS: 30 independently seeded 10,000-day paths "
    "under the baseline current-supervision regime, each discarding a "
    "2,000-day burn-in and exporting the daily market series and the primary "
    "listing's complete order-book event stream (about 270,000-300,000 events "
    "per seed). Every seed carries a provenance manifest recording the full "
    "configuration, independent market and supervision seeds, schema "
    "versions, and content hashes. Diagnostics are computed per seed under "
    "rules declared before execution, and cross-seed statements use the seed "
    "as the unit of inference; raw returns are never pooled across paths. An "
    "earlier illustrative diagnostic of the detached single-path "
    "simulation_results.csv is superseded by this campaign.")

VALIDATION_CAPTION = (
    "Table 8. Multi-seed financial stylized-facts diagnostics "
    "(30 independent seeds)")

VALIDATION_NOTE = (
    "Source: results/financial_validation/pilot_30seeds/diagnostics/. "
    "Statuses follow rules declared before execution; the seed is the unit "
    "of inference. R/P/N = seeds classified reproduced / partially "
    "reproduced / not reproduced. No simulation was executed during "
    "manuscript revision.")

VALIDATION_TABLE = [
    ["Fact", "Cross-seed diagnostic (means over 30 seeds)", "R/P/N",
     "Status"],
    ["Volatility clustering",
     "Absolute-return ACF 0.170 and squared-return ACF 0.183 at lag 1, "
     "beyond the 95% band at multiple lags",
     "30/0/0", "Reproduced"],
    ["Positive spread and depth",
     "Mean spread 1.38 ticks; two-sided depth in essentially all snapshots",
     "30/0/0", "Reproduced"],
    ["Trade-sign persistence",
     "Lag-1 trade-sign ACF 0.486 against an approximate 0.009 white-noise "
     "band",
     "30/0/0", "Reproduced"],
    ["Cancellation activity",
     "Cancellation-to-submission ratio 0.537",
     "30/0/0", "Reproduced"],
    ["Volume-volatility relation",
     "Mean daily volume-volatility correlation 0.113",
     "17/13/0", "Partially reproduced"],
    ["Fat-tailed returns",
     "Mean excess kurtosis 1.99 (cross-seed SD 1.33); mean Hill alpha 5.38",
     "3/27/0", "Partially reproduced"],
    ["Weak linear return autocorrelation",
     "Mean lag-1 return ACF -0.082, small but outside the white-noise band",
     "0/29/1", "Partially reproduced"],
    ["Positive price impact",
     "Mean correlation of log volume with immediate absolute mid-price "
     "movement -0.013 (95% half-width 0.004)",
     "0/4/26", "Not reproduced"],
]

VALIDATION_READING_1 = (
    "Four microstructure and volatility facts are reproduced in every seed. "
    "Two return-distribution facts are borderline. Excess kurtosis is "
    "positive in all seeds but exceeds 3 in only a minority, and Hill tail "
    "exponents average 5.4, so simulated tails are heavier than Gaussian yet "
    "thinner than the empirical 2-5 reference band in most seeds. Linear "
    "return autocorrelation is small but negative beyond the white-noise "
    "band at short lags, so the weak-autocorrelation criterion is only "
    "partially met. The multi-seed result also reverses the earlier "
    "single-path impression that fat tails were strong and linear "
    "predictability was severe; that detached path was unrepresentative, "
    "which is precisely why cross-seed validation is required.")

VALIDATION_READING_2 = (
    "Positive price impact is not reproduced. The implemented diagnostic "
    "correlates log trade volume with the immediate absolute mid-price "
    "movement; across seeds the mean correlation is -0.013 with a cross-seed "
    "95% half-width of 0.004, classified as not reproduced in 26 of 30 seeds "
    "and never as reproduced. Under this unsigned immediate-horizon measure, "
    "the engine's mid-price does not respond to trade size in the way "
    "empirical impact regularities require. A signed event-level diagnostic "
    "- trade sign multiplied by the future mid-price change over multiple "
    "event horizons, with seed-level uncertainty - is the sharper standard "
    "test and is identified as future work; under the pre-declared rule the "
    "fact is currently not reproduced and the paper reports it as such.")

VALIDATION_READING_3 = (
    "These results bound interpretation rather than invalidate the policy "
    "comparison. The policy outcomes in Sections 5.1-5.6 are claim-quality, "
    "greenhushing, procedural, and cost metrics generated by enforcement, "
    "information, and capacity mechanisms. As established in Section 3.7, "
    "firm environmental decisions contain no own-share-price feedback, so no "
    "policy ranking depends on price-mediated corporate incentives. What the "
    "partial validation does restrict is the reading of market-side "
    "quantities: investor signal distortion and price paths are internal "
    "model responses, and the financial market must not be described as "
    "empirically validated.")

LIMITATION_NINTH = (
    "Ninth, while the attached results complete all 200 parameter draws at "
    "all four horizons, only three stochastic replications are used per draw "
    "in the authoritative campaign. Paired differences reduce noise, but "
    "tail events and rare institutional failures may remain under-sampled. "
    "The executed 15-replication extension on 12 representative "
    "configurations narrows paired intervals in 92.3% of comparable cells "
    "yet confirms the three-replication default-weight winner in only 32 of "
    "48 draw-horizon cells, so draw-level winner labels are noisy and only "
    "distribution-level statements are emphasized. Partial rank correlations "
    "and standardized coefficients can also miss non-monotonicity and "
    "interactions.")

LIMITATION_ELEVENTH = (
    "Eleventh, the financial-market layer is partially validated and "
    "contains no market-discipline mechanism. Across 30 independent "
    "validation seeds, volatility clustering, positive spread and depth, "
    "trade-sign persistence, and cancellation activity are reproduced; fat "
    "tails, the volume-volatility relation, and weak linear return "
    "autocorrelation are only partially reproduced; and positive price "
    "impact is not reproduced under the implemented immediate-impact "
    "diagnostic. A static source audit additionally finds no own-share-price "
    "feedback into corporate greenwashing, disclosure, or "
    "transition-investment decisions; the only direct own-share market read "
    "is the treasury financing rule. The paper therefore claims no "
    "equity-price discipline of corporate environmental behavior, and the "
    "policy rankings do not depend on such a channel. Any conclusion that "
    "would require validated market microstructure or price-mediated "
    "corporate incentives is out of scope.")

DATA_AVAILABILITY = (
    "The complete model, parameter registry, tests, run-level outputs, "
    "manifests, figures, and reproducibility instructions are contained in "
    "the attached repository. The publication campaign is recorded in "
    "results/configuration.json, results/manifest.json, results/raw/, and "
    "results/summaries/. Strict inspection is read-only by default. The "
    "30-seed financial validation campaign, including per-seed provenance "
    "manifests, daily market series, order-book event streams, and "
    "multi-seed diagnostics, is in results/financial_validation/"
    "pilot_30seeds/. The executed 15-replication robustness extension is in "
    "results/replication_robustness/. Static audits of firm own-share price "
    "feedback and of paper-model claim consistency are in results/audits/. "
    "The earlier single-path diagnostic of the detached "
    "simulation_results.csv is retained only as provenance history and is "
    "superseded by the multi-seed campaign. The publication-readiness and "
    "legal audit is in docs/PUBLICATION_READINESS_REVIEW.md, and the claim "
    "classification is in docs/CLAIM_MATRIX.md. Latent-truth metrics are "
    "research-only and are never passed to agents during a run.")


def main() -> None:
    if not BACKUP.exists():
        shutil.copyfile(SOURCE, BACKUP)

    document = Document(str(SOURCE))

    # 1. Abstract.
    replace(find_paragraph(document, "Environmental-claim governance combines"),
            ABSTRACT)

    # 2. Section 3.7: explicit market-direction paragraph.
    insert_before(find_paragraph(document, "Employees update trust from"),
                  MARKET_DIRECTION)

    # 3. Section 4.4: executed robustness extension.
    heading = find_paragraph(
        document, "4.4 Replication robustness extension (not executed)")
    replace(heading, "4.4 Replication robustness extension")
    heading.style = "Heading 2"
    replace(find_paragraph(document, "Three paired replications per LHS draw"),
            ROBUSTNESS_DESIGN)
    replace(find_paragraph(document, "The extension has not been executed."),
            ROBUSTNESS_RESULT)

    # 4. Section 5.7: multi-seed financial validation.
    replace(find_paragraph(document, "A non-simulation diagnostic reads"),
            VALIDATION_INTRO)
    replace(find_paragraph(
        document, "Table 8. Preliminary financial stylized-facts"),
        VALIDATION_CAPTION)
    replace(find_paragraph(document, "Source: existing saved path only."),
            VALIDATION_NOTE)

    closing = find_paragraph(
        document, "A publication-grade validation should use pre-registered")
    old_table = None
    for table in document.tables:
        if table.cell(0, 0).text.strip() == "Fact":
            old_table = table
            break
    if old_table is None:
        raise ValueError("Table 8 (financial diagnostics) not found")
    insert_table_before(document, closing, VALIDATION_TABLE)
    old_table._tbl.getparent().remove(old_table._tbl)

    replace(closing, VALIDATION_READING_1)
    contributions = find_paragraph(document, "6. Contributions and Novelty")
    insert_before(contributions, VALIDATION_READING_2)
    insert_before(contributions, VALIDATION_READING_3)

    # 5. Section 5.5: investor distortion is market response, not discipline.
    frontier = find_paragraph(document, "The real-economy channels give these")
    frontier.add_run(
        " Because firm communication and investment rules contain no "
        "own-share-price feedback (Section 3.7), investor signal distortion "
        "is a market-response metric rather than evidence of price "
        "discipline.")

    # 6. Section 6: qualify the integration contribution.
    integration = find_paragraph(
        document, "The model's first contribution is institutional integration.")
    integration.add_run(
        " The financial market enters as an endogenous response and "
        "financing layer; the model does not implement, and the paper does "
        "not claim, own-share price discipline of corporate environmental "
        "choices.")

    # 7. Limitations: rewritten ninth, new eleventh.
    replace(find_paragraph(document, "Ninth, while the attached results"),
            LIMITATION_NINTH)
    insert_before(find_paragraph(document, "8. Conclusion"),
                  LIMITATION_ELEVENTH)

    # 8. Conclusion: future work on market microstructure and versioned
    #    market-discipline extension.
    conclusion = find_paragraph(
        document, "The model is best understood as a transparent computational "
        "laboratory")
    conclusion.add_run(
        " On the market side, future work should repair the non-reproduced "
        "price-impact diagnostic with signed event-level measures and "
        "seed-level inference, and, only if market discipline itself becomes "
        "the research question, introduce a pre-registered, information-safe "
        "own-share-price feedback channel as an explicitly versioned model "
        "extension.")

    # 9. Data and code availability.
    replace(find_paragraph(document, "The complete model, parameter registry"),
            DATA_AVAILABILITY)

    document.save(str(SOURCE))
    print(f"revised manuscript saved: {SOURCE}")
    print(f"pre-revision backup: {BACKUP}")


if __name__ == "__main__":
    main()
