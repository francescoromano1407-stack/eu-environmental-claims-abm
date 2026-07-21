"""Create the publication-readiness revision of the manuscript DOCX.

The original file is preserved.  This script performs targeted editorial
changes while retaining the existing document styles and embedded figures.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


SOURCE = Path("Francesco_Romano_EU_Environmental_Claims_ABM_Paper.docx")
OUTPUT = Path("Francesco_Romano_EU_Environmental_Claims_ABM_Paper_revised.docx")


def find_paragraph(document: Document, startswith: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(startswith):
            return paragraph
    raise ValueError(f"paragraph not found: {startswith!r}")


def replace(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def insert_before(reference: Paragraph, text: str,
                  style: str = "Normal") -> Paragraph:
    element = OxmlElement("w:p")
    reference._p.addprevious(element)
    paragraph = Paragraph(element, reference._parent)
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def insert_table_before(document: Document, reference: Paragraph,
                        rows: list[list[str]]) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        row_properties = table.rows[row_index]._tr.get_or_add_trPr()
        row_properties.append(OxmlElement("w:cantSplit"))
        if row_index == 0:
            row_properties.append(OxmlElement("w:tblHeader"))
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.text = value
            if row_index == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "1F4E78")
                cell._tc.get_or_add_tcPr().append(shading)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = None
                    color = OxmlElement("w:color")
                    color.set(qn("w:val"), "FFFFFF")
                    run._r.get_or_add_rPr().append(color)
    reference._p.addprevious(table._tbl)


def renumber_embedded_figure_titles(document: Document) -> None:
    """Keep titles printed inside reused campaign figures aligned with the paper."""
    titles = {
        "/word/media/image1.png":
            "Figure 1. Winner frequency by normative weight scenario",
        "/word/media/image2.png": "Figure 2. Horizon robustness",
    }
    font_path = Path("C:/Windows/Fonts/arial.ttf")
    for relationship in document.part.rels.values():
        if relationship.reltype != RT.IMAGE:
            continue
        image_part = relationship.target_part
        title = titles.get(str(image_part.partname))
        if title is None:
            continue
        image = Image.open(BytesIO(image_part.blob)).convert("RGB")
        draw = ImageDraw.Draw(image)
        draw.rectangle((0, 0, image.width, round(image.height * 0.073)),
                       fill="white")
        font = ImageFont.truetype(str(font_path),
                                  round(image.height * 0.037))
        box = draw.textbbox((0, 0), title, font=font)
        x = (image.width - (box[2] - box[0])) / 2
        draw.text((x, round(image.height * 0.024)), title,
                  fill="black", font=font)
        output = BytesIO()
        image.save(output, format="PNG")
        image_part._blob = output.getvalue()


def main() -> None:
    document = Document(SOURCE)

    replace(find_paragraph(document, "Environmental claims can influence"),
            "Environmental-claim governance combines an information problem "
            "with an institutional-capacity problem: regulators observe claims "
            "and uncertain evidence rather than firms' latent environmental "
            "states, while preventive scrutiny may suppress truthful "
            "communication. We develop an information-safe agent-based model "
            "linking corporate environmental claims, consumer demand, investor "
            "valuation, workforce responses, a limit-order-book market, and "
            "capacity-constrained supervision. We compare current EU-style "
            "ex-post supervision with two experimental institutions: a voluntary "
            "algorithmic SME pre-screening hub and a certified green-data "
            "connector. A 28-parameter Latin-hypercube design evaluates 200 "
            "configurations at 120, 365, 1,000, and 2,000 days, with three paired "
            "common-random-number replications per configuration. Within the "
            "model, the hub reduces exposure-weighted claim severity in "
            "85.0-90.5% of configurations across horizons, but systematically "
            "increases greenhushing and public and firm compliance costs. The "
            "connector is weak and heterogeneous at 120 days, improves severity "
            "in 60.5% of configurations at 365 days, and becomes robust at 1,000 "
            "and 2,000 days (97.0% and 100.0%), while remaining cost-intensive "
            "and increasing evidence-conflict workload. Default rankings reverse "
            "across horizons in 80.5% of matched configurations, showing why "
            "policy comparisons must be conditional on time, evidence quality, "
            "capacity, participation, and normative weights. A preliminary "
            "diagnostic of one saved financial path reproduces fat tails, partly "
            "reproduces volatility clustering, and fails the weak-return-"
            "autocorrelation criterion; financial validation therefore remains "
            "incomplete. The results are transparent simulation-based policy "
            "orderings, not empirical forecasts or causal estimates.")

    replace(find_paragraph(document, "The paper makes six contributions."),
            "The paper contributes a computational institutional-design "
            "framework for environmental-claim supervision. Its novelty lies in "
            "combining three elements that are usually modeled separately: an "
            "information-safe architecture in which no decision-maker observes "
            "latent environmental truth; procedural enforcement with queues, "
            "correction windows, assurance, track-gated remedies, and capacity-"
            "consuming evidence conflict; and endogenous feedback from claims "
            "and public decisions into consumption, finance, and workforce "
            "outcomes. This integration exposes two conditional mechanisms. "
            "Preventive pre-screening trades lower modeled overstatement for more "
            "greenhushing and higher cost, whereas certified data infrastructure "
            "produces delayed claim-quality benefits that become robust only over "
            "longer horizons. The paired global design maps where these mechanisms "
            "hold and why normative rankings change; it does not estimate real-"
            "world policy effects.")

    replace(find_paragraph(document, "Directive (EU) 2019/2161 strengthens"),
            "Directive (EU) 2019/2161 strengthens enforcement and penalty "
            "availability for widespread infringements and widespread "
            "infringements with a Union dimension (European Parliament and "
            "Council, 2019). It requires Member States to set the maximum fine "
            "for the relevant infringements at no less than 4% of the trader's "
            "annual turnover in the Member State or Member States concerned. It "
            "does not establish 4% as a uniform EU ceiling. The unchanged model "
            "therefore treats an exact 4% cap as a LEGAL-ANCHOR scenario choice, "
            "available only when a case is classified as a coordinated widespread "
            "cross-border consumer infringement. Ordinary simulated cases use a "
            "separate 1% experimental cap applied to a simulation-scale turnover "
            "proxy.")

    replace(find_paragraph(document, "The CSDDD is represented only"),
            "The CSDDD is represented only as a reduced-form due-diligence "
            "procedure (European Parliament and Council, 2024b, as amended in "
            "2026). Its 3% maximum is callable only for due-diligence cases, and "
            "the model's application date is 26 July 2029. Directive (EU) "
            "2026/470 also revises the principal undertaking threshold to more "
            "than 5,000 employees and more than EUR 1.5 billion net worldwide "
            "turnover, with separate group, franchise, and third-country rules. "
            "The current model gates CSDDD by date and remedy track but does not "
            "implement that complete firm-scope test; it must therefore not be "
            "described as a doctrinally complete CSDDD implementation. "
            "Greenhushing is not modeled as an offence.")

    model_heading = find_paragraph(document, "3. Model and Methodology")
    insert_before(model_heading, "2.3 Related work and positioning", "Heading 2")
    insert_before(model_heading,
                  "Agent-based compliance research shows how inspection, social "
                  "interaction, learning, and enforcement can generate aggregate "
                  "reporting behavior (Korobow, Johnson, and Axtell, 2007). "
                  "Recent inspection models also combine behavioral theory and "
                  "administrative data. These studies motivate endogenous "
                  "compliance but generally do not couple uncertain claim "
                  "evidence, sequenced legal remedies, greenhushing, and a "
                  "financial market.")
    insert_before(model_heading,
                  "The RegTech and SupTech literature emphasizes automated data "
                  "collection and analysis, prospective surveillance, resource "
                  "constraints, data quality, explainability, and operational "
                  "risk (Broeders and Prenio, 2018; Financial Stability Board, "
                  "2020). The experimental hub and connector instantiate these "
                  "institutional tensions; they are not implementations of "
                  "existing EU obligations.")
    insert_before(model_heading,
                  "Financial agent-based models are commonly assessed against "
                  "fat-tailed returns, volatility clustering, and limited linear "
                  "return predictability (Cont, 2001; Lux and Marchesi, 1999). "
                  "Reproducing selected stylized facts is weaker than empirical "
                  "calibration or out-of-sample validation. This paper therefore "
                  "reports each diagnostic as reproduced, partially reproduced, "
                  "or not reproduced.")
    insert_before(model_heading,
                  "The paper's defensible novelty claim concerns integration. To "
                  "our knowledge, it is among the first computational policy "
                  "experiments to combine information-safe environmental-claim "
                  "supervision, explicit legal procedure and evidentiary "
                  "conflict, and endogenous financial and real-economy feedback "
                  "in one paired institutional comparison. No absolute 'first' "
                  "claim is made without a systematic review.")

    replace(find_paragraph(document, "The reproducibility document also preserves"),
            "The authoritative design is the machine-readable configuration in "
            "results/configuration.json: all four horizons have subset: null and "
            "all four summaries contain 200 draws. The current reproducibility "
            "guide and data dictionary have been synchronized to this layout. "
            "Historical selective-confirmation commands are not part of the "
            "reported design.")

    findings_heading = find_paragraph(document, "5. Findings")
    insert_before(findings_heading,
                  "4.4 Replication robustness extension (not executed)",
                  "Heading 2")
    insert_before(findings_heading,
                  "Three paired replications per LHS draw are appropriate for "
                  "broad global screening but can under-sample rare escalation, "
                  "institutional failure, and tail outcomes. A separate dry-run-"
                  "first workflow deterministically selects 12 configurations by "
                  "maximin coverage of standardized cross-horizon outcomes and "
                  "targets 15 paired replications by default (configurable from "
                  "10 to 20). It preserves the authoritative first three "
                  "replications and would run only replications 4-15 under the "
                  "same paired seed schedule.")
    insert_before(findings_heading,
                  "The extension has not been executed. Its selected draws are "
                  "41, 42, 69, 71, 75, 104, 109, 136, 144, 147, 180, and 187 at "
                  "each horizon, and every extension file is currently pending. "
                  "No interval-narrowing or ranking-stability result from this "
                  "workflow is reported in the paper.")

    replace(find_paragraph(document, "5.1 No universal winner"),
            "5.1 Conditional policy maps")
    replace(find_paragraph(document, "The central result is ranking instability."),
            "Ranking instability explains why the two policy mechanisms must be "
            "reported conditionally rather than collapsed into a universal "
            "league table. The default-weight winner changes across horizons for "
            "161 of 200 matched parameter draws (80.5%). Within a horizon, "
            "changing normative weights changes the winner in 70.0% of 120-day "
            "draws, 63.0% of 365-day draws, 58.0% of 1,000-day draws, and 45.5% "
            "of 2,000-day draws. These reversals are consequences of distinct "
            "accuracy, greenhushing, cost, and timing trade-offs, not evidence "
            "that comparison is impossible.")

    replace(find_paragraph(document, "Figure 9. Winner frequency"),
            "Figure 1. Winner frequency by normative weight scenario at 120 days.")
    replace(find_paragraph(document, "Figure 10. Default-weight"),
            "Figure 2. Default-weight winner shares across horizons.")

    contributions = find_paragraph(document, "6. Contributions and Novelty")
    insert_before(contributions, "5.7 Financial-market validation", "Heading 2")
    insert_before(contributions,
                  "A non-simulation diagnostic reads the existing 2,000-day "
                  "simulation_results.csv path. The file contains 2,001 prices. "
                  "Exactly zero calendar returns account for 28.5%, so the "
                  "primary tests use 1,430 non-zero price-change returns and "
                  "record that filter. Because the CSV has no matching seed and "
                  "configuration manifest, this is an illustrative single-path "
                  "audit rather than publication-grade validation.")
    insert_before(contributions,
                  "Table 8. Preliminary financial stylized-facts diagnostics",
                  "Caption")
    insert_before(contributions,
                  "Source: existing saved path only. Transparent rules and "
                  "provenance hash are in results/validation/. No model run was "
                  "executed to produce this table.", "Table Note")
    insert_table_before(document, contributions, [
        ["Fact", "Diagnostic", "Status", "Interpretation"],
        ["Fat-tailed returns",
         "Excess kurtosis 21.66; Hill alpha 3.12; Jarque-Bera p effectively 0",
         "Reproduced", "Strong on this path; not cross-seed validation"],
        ["Volatility clustering",
         "Absolute-return ACF 0.327 at lag 1 and 0.071 at lag 5; squared-return Ljung-Box(20) p≈4.0e-42",
         "Partially reproduced", "Short-lag persistence fades by lags 10-20"],
        ["Weak linear return autocorrelation",
         "Return ACF 0.445 at lag 1, 0.287 at lag 2, 0.170 at lag 5; Ljung-Box(20) p≈3.4e-129",
         "Not reproduced", "The saved path is too linearly predictable"],
        ["Order-book microstructure",
         "Spread, depth, volume, cancellations, trade signs, and impact not exported",
         "Pending", "Cannot be inferred from daily price alone"],
    ])
    insert_before(contributions,
                  "A publication-grade validation should use pre-registered, "
                  "independently seeded paths at fixed baseline parameters and "
                  "export order-level spread, depth, volume, cancellation, trade-"
                  "sign, and price-impact series. The current result supports fat "
                  "tails, gives limited support to volatility persistence, and "
                  "rejects the weak-linear-autocorrelation criterion for this "
                  "path. The manuscript therefore does not describe the financial "
                  "market as empirically validated.")

    replace(find_paragraph(document, "Ninth, while the attached results complete"),
            "Ninth, while the attached results complete all 200 parameter draws "
            "at all four horizons, only three stochastic replications are used "
            "per draw. Paired differences reduce noise, but tail events and rare "
            "institutional failures may remain under-sampled. The separate 10-20 "
            "replication workflow is implemented but unexecuted, so it supplies "
            "no robustness result yet. Partial rank correlations and standardized "
            "coefficients can also miss non-monotonicity and interactions.")
    replace(find_paragraph(document, "Tenth, the documentation contains"),
            "Tenth, legal and reproducibility fidelity remain incomplete. The "
            "CSDDD layer lacks the amended undertaking-scope gate, several market "
            "and sustainable-finance tracks are reduced-form, and the shipped "
            "campaign manifests record a git hash suffixed +dirty. The archived "
            "release must preserve the exact dirty-tree diff or reproduce results "
            "from a clean tagged tree. These limitations do not authorize claims "
            "of legal completeness or bit-level reproducibility.")

    replace(find_paragraph(document, "This study develops an information-safe"),
            "This study develops an information-safe agent-based model for "
            "comparing three institutional approaches to environmental-claim "
            "governance in the European Union. Its main result is mechanism-"
            "based. Algorithmic pre-screening frequently lowers material "
            "overstatement and exposure-weighted severity, but it does so with a "
            "persistent greenhushing and public/firm cost penalty. Certified "
            "environmental-data infrastructure produces weak or heterogeneous "
            "short-run effects and much stronger long-horizon claim-quality "
            "benefits, while remaining expensive and loading the evidence-"
            "conflict process.")
    replace(find_paragraph(document, "No regime is universally superior."),
            "Ranking instability is evidence for conditional policy analysis, "
            "not the paper's endpoint. Default rankings reverse across horizons "
            "in 80.5% of shared draws because the institutions differ in timing, "
            "accuracy, greenhushing, and cost. The defensible conclusion is that "
            "institutional performance depends on horizon, participation, "
            "evidence quality, administrative capacity, discounting, and "
            "objectives; the model does not identify a real-world optimal policy.")

    replace(find_paragraph(document, "The complete model, parameter registry"),
            "The complete model, parameter registry, tests, run-level outputs, "
            "manifests, figures, and reproducibility instructions are contained "
            "in the attached repository. The publication campaign is recorded in "
            "results/configuration.json, results/manifest.json, results/raw/, and "
            "results/summaries/. Strict inspection is read-only by default. "
            "Financial diagnostics and their source hash are in "
            "results/validation/. The unexecuted replication extension is "
            "implemented in run_replication_robustness.py. The publication-"
            "readiness and legal audit is in "
            "docs/PUBLICATION_READINESS_REVIEW.md. Latent-truth metrics are "
            "research-only and are never passed to agents during a run.")

    replace(find_paragraph(document, "Table 8. Parameter-role taxonomy"),
            "Table 9. Parameter-role taxonomy and sampled families")

    appendix = find_paragraph(document, "Appendix A. Parameter Classification")
    new_references = [
        "Broeders, D., and Prenio, J. (2018). Innovative technology in financial supervision (suptech): The experience of early users. FSI Insights No. 9, Bank for International Settlements. https://www.bis.org/fsi/publ/insights9.htm",
        "Cont, R. (2001). Empirical properties of asset returns: Stylized facts and statistical issues. Quantitative Finance, 1(2), 223-236. https://doi.org/10.1080/713665670",
        "Financial Stability Board. (2020). The use of supervisory and regulatory technology by authorities and regulated institutions. https://www.fsb.org/2020/10/fsb-report-highlights-increased-use-of-regtech-and-suptech/",
        "Korobow, A., Johnson, C., and Axtell, R. L. (2007). An agent-based model of tax compliance with social networks. National Tax Journal, 60(3), 589-610. https://doi.org/10.17310/ntj.2007.3.16",
        "Lux, T., and Marchesi, M. (1999). Scaling and criticality in a stochastic multi-agent model of a financial market. Nature, 397, 498-500. https://doi.org/10.1038/17290",
    ]
    for reference in new_references:
        insert_before(appendix, reference, "Reference")
    references = [paragraph for paragraph in document.paragraphs
                  if paragraph.style.name == "Reference"]
    for paragraph in sorted(references, key=lambda item: item.text.casefold()):
        appendix._p.addprevious(paragraph._p)

    # Legal-map table: make the CSDDD implementation boundary explicit.
    legal_table = document.tables[0]
    legal_table.cell(6, 1).text = (
        "Due-diligence date/remedy gate; full undertaking scope not implemented")
    legal_table.cell(6, 3).text = (
        "Reduced-form binding-law representation requiring legal review")

    document.core_properties.title = (
        "Preventing Environmental-Claim Misrepresentation in the European Union")
    document.core_properties.subject = (
        "Publication-readiness revision; simulation-based policy experiment")
    renumber_embedded_figure_titles(document)
    document.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    main()
